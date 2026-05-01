"""Document ingestion pipeline.

Orchestrates all six stages of document processing:

  Stage 1: Format detection (extension + magic bytes)
  Stage 2: Content extraction (text + page/heading markers)
  Stage 3: Structure extraction (domain-specific section hierarchy)
  Stage 4: OpenViking hierarchy builder (viking:// directory structure)
  Stage 5: L0/L1/L2 generation (summaries written to OpenViking)
  Stage 6: Qdrant dual-write (embeddings with full payload)

Returns a :class:`~app.naturalsentinel.memory.pg_models.PgDocument` instance
(not yet committed to the DB — the caller commits).
"""

from __future__ import annotations

import base64
import logging
import mimetypes
import os
import uuid
from datetime import UTC, datetime
from typing import Any

# OV write helpers live in openviking_service (extracted in Phase R).
# Re-exported below so legacy imports like
# ``from app.naturalsentinel.documents.pipeline import build_openviking_hierarchy``
# keep working until the next major release.
from app.naturalsentinel.documents.constants import OV_DOCUMENT_ROOT
from app.naturalsentinel.documents.openviking_service import build_openviking_hierarchy

logger = logging.getLogger(__name__)

__all__ = ["build_openviking_hierarchy", "ingest_document"]


# ---------------------------------------------------------------------------
# Stage 1: Format detection
# ---------------------------------------------------------------------------

_MAGIC_BYTES: list[tuple[bytes, str]] = [
    (b"%PDF", "application/pdf"),
    (
        b"PK\x03\x04",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    (b"<html", "text/html"),
    (b"<!DOCTYPE", "text/html"),
    (b"# ", "text/markdown"),
    (b"---\n", "text/markdown"),
]


def detect_format(file_name: str, content: bytes) -> str:
    """Return MIME type for a document.

    Checks magic bytes first, then falls back to filename extension.
    """
    for magic, mime in _MAGIC_BYTES:
        if content[: len(magic)].lower().startswith(magic.lower()):
            return mime
    mime, _ = mimetypes.guess_type(file_name)
    return mime or "text/plain"


# ---------------------------------------------------------------------------
# Stage 2: Content extraction
# ---------------------------------------------------------------------------


def extract_text(content: bytes, mime_type: str, file_name: str) -> str:
    """Extract plain text from raw file bytes.

    Uses pdfplumber for PDF, python-docx for DOCX, and simple UTF-8
    decoding for HTML/Markdown/text. Falls back to UTF-8 with replacement.
    """
    if mime_type == "application/pdf":
        return _extract_pdf(content)
    if mime_type.endswith("wordprocessingml.document"):
        return _extract_docx(content)
    if mime_type == "text/html":
        return _extract_html(content)
    # Markdown, plain text, or unknown
    return content.decode("utf-8", errors="replace")


def _extract_pdf(content: bytes) -> str:
    try:
        import io

        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text:
                    pages.append(f"[Page {page.page_number}]\n{text}")
            return "\n\n".join(pages)
    except ImportError:
        logger.warning("pdfplumber not installed; treating PDF as binary (no text)")
        return "[PDF content — pdfplumber not installed]"
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        return ""


def _extract_docx(content: bytes) -> str:
    try:
        import io

        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed; cannot extract DOCX text")
        return "[DOCX content — python-docx not installed]"
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""


def _extract_html(content: bytes) -> str:
    from app.naturalsentinel.fetchers.live.parsers import html_to_text

    html = content.decode("utf-8", errors="replace")
    return html_to_text(html)


# ---------------------------------------------------------------------------
# Stage 3: Structure extraction (delegated to extractors module)
# ---------------------------------------------------------------------------


def build_document_tree(
    raw_text: str,
    doc_id: str,
    doc_type: str,
    source_url: str,
    file_name: str,
    file_size: int,
    metadata: dict | None = None,
):
    from app.naturalsentinel.documents.extractors import extract_structure

    return extract_structure(
        raw_text=raw_text,
        doc_id=doc_id,
        doc_type=doc_type if doc_type else "auto",
        source_url=source_url,
        file_name=file_name,
        file_size=file_size,
        metadata=metadata,
    )


# ---------------------------------------------------------------------------
# Stage 4 + 5: OpenViking hierarchy
# ---------------------------------------------------------------------------
# See openviking_service.build_openviking_hierarchy — imported at the top
# of this module and re-exported for legacy callers.


# ---------------------------------------------------------------------------
# Stage 6: Qdrant dual-write
# ---------------------------------------------------------------------------


def write_sections_to_qdrant(
    qdrant_client, tree, root_viking_uri: str, tags: list[str]
) -> int:
    """Write all document sections to Qdrant.

    Returns the number of points upserted.
    """
    from app.naturalsentinel.documents.qdrant_service import upsert_document_sections

    sections: list[dict] = []
    for node in tree.all_nodes():
        node_uri = f"{root_viking_uri}/{node.uri_path}"
        chunk_id = f"{tree.doc_id}:{node.uri_path}:0"
        sections.append(
            {
                "chunk_id": chunk_id,
                "viking_uri": node_uri,
                "section_path": node.section_path,
                "node_type": node.node_type,
                "title": node.title,
                "text": node.text,
                "abstract": node.abstract,
                "overview": node.overview,
                "line_start": node.line_start,
                "line_end": node.line_end,
                "char_offset_start": node.char_offset_start,
                "char_offset_end": node.char_offset_end,
                "page_number": node.page_number,
                "word_count": node.word_count(),
            }
        )

    return upsert_document_sections(
        client=qdrant_client,
        doc_id=tree.doc_id,
        doc_type=tree.doc_type,
        source_url=tree.source_url,
        tags=tags,
        sections=sections,
    )


# ---------------------------------------------------------------------------
# Main ingest entry point
# ---------------------------------------------------------------------------


def ingest_document(
    *,
    source_url: str = "",
    file_path: str = "",
    content_b64: str = "",
    content_type: str = "",
    doc_type: str = "auto",
    metadata: dict[str, Any] | None = None,
    ov_client=None,
    qdrant_client=None,
    session_db=None,
) -> dict[str, Any]:
    """Full ingestion pipeline for a single document.

    Exactly one of ``source_url``, ``file_path``, or ``content_b64`` must be
    provided.  Returns a dict suitable for creating a PgDocument record.

    Args:
        source_url: Fetch document from this URL.
        file_path: Read document from local filesystem path.
        content_b64: Base64-encoded file content (agent upload).
        content_type: MIME type hint for base64 content.
        doc_type: Document type hint. ``"auto"`` triggers auto-detection.
        metadata: Optional extra metadata (client_name, matter_id, tags, etc.).
        ov_client: OpenViking SyncHTTPClient. May be None (hierarchy step skipped).
        qdrant_client: Qdrant client. May be None (embedding step skipped).
        session_db: SQLModel Session for PgDocument CRUD. May be None.

    Returns:
        Dict with doc_id, viking_uri, title, doc_type, section_count, status,
        structure_summary.
    """
    doc_id = str(uuid.uuid4())
    meta = metadata or {}
    tags: list[str] = meta.get("tags", [])

    # ── Stage 1+2: Fetch + extract text ───────────────────────────────────
    if source_url:
        raw_content, file_name = _fetch_url(source_url)
        effective_content_type = content_type or detect_format(file_name, raw_content)
    elif file_path:
        if not os.path.isabs(file_path):
            raise ValueError("file_path must be an absolute path")
        file_name = os.path.basename(file_path)
        with open(file_path, "rb") as fh:
            raw_content = fh.read()
        effective_content_type = content_type or detect_format(file_name, raw_content)
    elif content_b64:
        raw_content = base64.b64decode(content_b64)
        file_name = meta.get("file_name", "document.bin")
        effective_content_type = content_type or detect_format(file_name, raw_content)
    else:
        raise ValueError("One of source_url, file_path, or content_b64 is required")

    file_size = len(raw_content)
    raw_text = extract_text(raw_content, effective_content_type, file_name)

    if not raw_text.strip():
        logger.warning("Document %s yielded empty text — check parser", file_name)

    if source_url:
        effective_source = source_url
    elif file_path:
        effective_source = f"file://{file_path}"
    else:
        effective_source = f"upload://{file_name}"

    # ── Stage 3: Structure extraction ─────────────────────────────────────
    tree = build_document_tree(
        raw_text=raw_text,
        doc_id=doc_id,
        doc_type=doc_type,
        source_url=effective_source,
        file_name=file_name,
        file_size=file_size,
        metadata=meta,
    )

    # ── Stage 4+5: OpenViking hierarchy ───────────────────────────────────
    viking_uri = f"{OV_DOCUMENT_ROOT}/{doc_id}"
    if ov_client is not None:
        try:
            viking_uri = build_openviking_hierarchy(ov_client, tree)
        except Exception as exc:
            logger.warning("OpenViking hierarchy build failed: %s", exc)

    # ── Stage 6: Qdrant dual-write ─────────────────────────────────────────
    qdrant_point_count = 0
    if qdrant_client is not None:
        try:
            qdrant_point_count = write_sections_to_qdrant(
                qdrant_client, tree, viking_uri, tags
            )
        except Exception as exc:
            logger.warning("Qdrant write failed: %s", exc)

    # ── Persist PgDocument record ──────────────────────────────────────────
    status = "ready"
    if session_db is not None:
        try:
            _persist_pg_document(
                session=session_db,
                doc_id=doc_id,
                tree=tree,
                viking_uri=viking_uri,
                meta=meta,
            )
        except Exception as exc:
            logger.warning("PgDocument persist failed: %s", exc)
            status = "ready"  # non-fatal

    structure_summary = (
        tree.root_nodes[0].abstract if tree.root_nodes else raw_text[:200]
    )

    return {
        "doc_id": doc_id,
        "uri": viking_uri,
        "title": tree.title,
        "doc_type": tree.doc_type,
        "file_name": tree.file_name,
        "file_size": tree.file_size,
        "source_url": tree.source_url,
        "section_count": tree.section_count(),
        "qdrant_points": qdrant_point_count,
        "status": status,
        "structure_summary": structure_summary,
        "metadata": meta,
    }


def _fetch_url(url: str) -> tuple[bytes, str]:
    """Fetch a URL and return (content_bytes, file_name)."""
    import urllib.request

    req = urllib.request.Request(
        url,
        headers={"User-Agent": "NaturalSentinel/1.0 document-ingest"},
    )
    with urllib.request.urlopen(req, timeout=60) as resp:
        content = resp.read()
    # Derive filename from URL path
    path_part = url.split("?")[0].rstrip("/")
    file_name = path_part.split("/")[-1] or "document"
    if "." not in file_name:
        ct = resp.headers.get("Content-Type", "")
        if "pdf" in ct:
            file_name += ".pdf"
        elif "html" in ct:
            file_name += ".html"
        else:
            file_name += ".txt"
    return content, file_name


def _persist_pg_document(
    *, session, doc_id: str, tree, viking_uri: str, meta: dict
) -> None:
    """Persist a PgDocument record to PostgreSQL.

    Populates source_url / domain / jurisdiction as first-class columns
    (Phase P0.4) in addition to the metadata_json blob — the columns
    are the canonical query path, the blob is redundant alignment with
    Qdrant / OV payloads.
    """
    from app.naturalsentinel.memory.pg_models import PgDocument

    now = datetime.now(UTC)
    structure = {"nodes": [n.uri_path for n in tree.root_nodes]}

    persisted_meta = dict(meta)
    persisted_meta["source_url"] = tree.source_url

    db_obj = PgDocument(
        doc_id=doc_id,
        title=tree.title,
        doc_type=tree.doc_type,
        file_name=tree.file_name,
        file_size=tree.file_size,
        viking_uri=viking_uri,
        section_count=tree.section_count(),
        status="ready",
        source_url=tree.source_url or "",
        domain=meta.get("domain"),
        jurisdiction=meta.get("jurisdiction", "federal"),
        metadata_json=persisted_meta,
        structure_json=structure,
        created_at=now,
        updated_at=now,
        created_by=meta.get("created_by", ""),
    )
    session.add(db_obj)
    session.commit()
    session.refresh(db_obj)
